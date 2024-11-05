#!/usr/bin/python3
# -*- coding: utf-8 -*-

# автор: Титовский С.А.

import wx
import wx.lib.scrolledpanel
import sys
import requests
import re
import time
import string # for buttons
import datetime  # импорт библиотеки дат и времени
from datetime import datetime, timedelta
import os  # импорт библиотеки для работы с операционной системой
import docx # импорт библиотек для работы с Word
from docx import Document # импорт библиотек для работы с Word
from docx.enum.style import WD_STYLE_TYPE # импорт библиотек для работы с Word
from docx.enum.text import WD_ALIGN_PARAGRAPH # импорт библиотек для работы с Word
from docx.shared import Pt, Cm, Inches # импорт библиотек для работы с Word
from docx.oxml.table import CT_Row, CT_Tc # импорт библиотек для работы с Word
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement # импорт библиотек для работы с Word
from docx.oxml.ns import qn # импорт библиотек для работы с Word
import itertools # инструмент для итерации (в данном случае для формирования первой строки таблицы)
from itertools import islice # инструмент для итерации (в данном случае для формирования первой строки таблицы)
import socket
import sqlite3
#import openpyxl # импорт библиотек для работы Excel
#from openpyxl import Workbook # импорт библиотек для работы Excel
import locale # для работы с локалью
import threading
from threading import Thread # для работы с потоками
from pubsub import pub # для работы с подписчиками
import ftplib # для работы с ftp
import queue
import pygame
from pygame.locals import *

#=========================================
#=========================================
#=========================================
#=========================================        
# Thread of making List
class ListThread(Thread):
    
    def __init__(self, arg, UserDate, DocDir):
        super().__init__()
        self.i = arg
        self.UserDate = UserDate

    def run(self):

        # Спрячем основное окно
        WinOsn.Show(False)

        UpdDlg = wx.GenericProgressDialog(
            "Составление списка присутствующих на совещание\n"  + str(spisoksov[5][self.i]),
            "Составление списка присутствующих на совещание\n"  + str(spisoksov[5][self.i]),
            maximum = 16,
            parent = None,
            #style = wx.PD_AUTO_HIDE|wx.PD_SMOOTH)
            style = wx.PD_AUTO_HIDE|wx.PD_APP_MODAL|wx.PD_SMOOTH)

        time.sleep(0.5)

        try:
            # выполнение функции создания справки
            formfile(spisoksov[2][self.i], self.UserDate, spisoksov[4][self.i], spisoksov[3][self.i], DocDir = DocDir, Dlg = UpdDlg)

            # отправка обновения основному окну
            UpdDlg.Update(UpdDlg.GetRange() - 1, "Успешное выполнения выгрузки справки")
            time.sleep(0.5)
            UpdDlg.Destroy()

        except Exception as Err:
            ToLog("Ошибка при выполнении функции справки, ошибка = " + str(Err))
            #raise Exception

        # Покажем основное окно
        WinOsn.Show(True)

#=========================================
#=========================================
#=========================================
#=========================================        
# OsnFrame
class OsnWindow(wx.Frame):

    def __init__(self, parent, spisoksov, WinPos = wx.DefaultPosition, DateUser = "01.01.2021", HSize = False, DocDir = os.getcwd()):
        wx.Frame.__init__(
            self, parent, -1, "Список совещаний")
        
        self.WinPos = WinPos
        self.Date = DateUser
        self.spisok = spisoksov
        self.Difference = False
        self.temp = []
        self.HSize = HSize
        self.DocDir = DocDir

        frameIcon = wx.Icon(os.getcwd() + "\\images\\IconPNG.png")
        self.SetIcon(frameIcon)

        pygame.init()
        pygame.mixer.init()

        global Region
        #countting users if it's Yprava
        if Region == 0:
            self.Updating = wx.GenericProgressDialog(
                " ","Подождите, идёт подсчет участников совещаний...",
                maximum = len(self.spisok), parent = self,
                style = wx.PD_AUTO_HIDE|wx.PD_APP_MODAL|wx.PD_SMOOTH)
            self.FactSpisok = self.CountFact()
            self.Updating.Destroy()
            self.OpenPanel()

        else:
            self.FactSpisok = False
            self.OpenPanel()

        self.Bind(wx.EVT_CLOSE, self.OnCloseWindow)
        ToLog("Main Frame opened on Date = " + str(self.Date))
#=========================================
    def OnCloseWindow(self, evt):
        #print("Нажата кнопка Закрыть, осуществляю выход из программы")
        try:
            ToLog("Application closed by User's command")
            self.Show(False)
            try:
                self.panel.MonitFrame.Destroy()
            except Exception:
                pass

            ToLog("Stopping LogThread")
            global threadLog
            threadLog.stop = True
            threadLog.join()
            #self.StopThread()
            evt.Skip()
            sys.exit()
        except Exception as Err:
            ToLog("Error in OnClose Window< error code = " + str(Err))
            evt.Skip()
            #sys.exit()
            #raise Exception
         
#=========================================
    def CountFact(self):
        global Region
        SpisFact = []
        SpisSetka = []
        Label = []
        SpisRazn = []

        for i in range (0, len(self.spisok[5])):
            # подсчет приглашенных по всем совещаниям
            SpisFact.append(self.countFact(self.spisok[2][i]))

        return SpisFact
#================================================
    def countFact(self, idsov): 
        itogitogov = []
        ssilkaSS = str(
            "http://10.132.71.156/pls/ss/selector.report.study_p?sid=" + idsov +
            "&us=" + str(Region))
        # Запрашиваем код
        try:
            responseSS = requests.get(ssilkaSS)
        except Exception as Err:
            ToLog("Error in countFact, Error code = " + str(Err))
            sys.exit()
            
        filesplit = responseSS.text.splitlines()    
        dolgnost = []
        
        #Обработка кода страницы и составление списков
        for i in range(0, len(filesplit)-1):
               
            # Формируем списки для таблицы - должность, ФИО, Примечание

            if (
                filesplit[i].find('''<tr><td colspan=3 class=z2>''')!=-1
                or
                filesplit[i].find('''<tr><td class=spr valign=top>''')!=-1):
                
                if filesplit[i].find('''<tr><td colspan=3 class=z2>''')!=-1:
                    dolgnost.append("КАБИНЕТ" + str(filesplit[i][filesplit[i].find('''z2>''')+3:filesplit[i].find('''</td>''')]))

                if filesplit[i].find('''<tr><td class=spr valign=top>''')!=-1:
                    dolgnost.append(filesplit[i+1])

        # Убираем кабинеты и студии из списка
    
        dolgnostFiltr = []
    
        for i in range(0, len(dolgnost)):
            if dolgnost[i].find("КАБИНЕТ") == -1:
                dolgnostFiltr.append(dolgnost[i])
            
                #print ("Долж = "+dolgnost[i])
            #print ("Долж = "+dolgnost[i]+" --- ФИО = "+fio1[i]+" --- Прим = "+prim[i])
   
        return dolgnostFiltr

#=========================================
    def OpenPanel(self):
        #form OsnPanel
        self.panel = OsnPanel(self, self.spisok, self.Date, self.FactSpisok, self.DocDir)
        try:
            self.Move(self.WinPos)
        except Exception as Err:
            SomeError(None, "Failed to move Control Window to saved position")
            self.Centre()

        self.panel.SetSizer(self.panel.CommonVbox)
        self.Layout()
        if self.HSize == False:
            self.HSize = 320
        self.SetClientSize((self.HSize, len(self.spisok[5]) * 30 + 250))
        self.Show(True)

#=========================================
#=========================================
#=========================================
#=========================================
#OsnPanel
class OsnPanel(wx.lib.scrolledpanel.ScrolledPanel):
#class OsnPanel(wx.Panel):

    def __init__(self, Parent, spisok, date, FactList, DocDir):
        wx.lib.scrolledpanel.ScrolledPanel.__init__(self, parent = Parent)
        self.SetupScrolling()
        #wx.Panel.__init__(self, parent = Parent)
        global MyDir
        self.spisok = spisok
        self.Date = date
        self.spisoksov = spisok
        self.frame = Parent
        self.FactList = FactList
        self.DocDir = DocDir
        #print("List on Panel")
        #for List in self.spisoksov:
        #    print(str(List))
        #print("FactList")
        #print(str(self.FactList))

        #read colours (and TimeRefresh)from file
        self.colour = LoadColours(self.DocDir)
        
        self.SetBackgroundColour(wx.Colour(self.colour[0]))
        self.CommonVbox = wx.FlexGridSizer(rows = len(spisoksov[5]) + 4, cols = 1, hgap = 6, vgap = 6)
        self.CommonVbox.AddGrowableCol(0, 1)
        for rows in range (0, len(spisoksov[5]) + 4):
            self.CommonVbox.AddGrowableRow(rows, 0)
        ButtonVbox = wx.FlexGridSizer(rows = 1, cols = 3, hgap = 6, vgap = 6)
        ButtonVbox.AddGrowableCol(1, 1)

        # Добавление кнопки сброса
        self.buttonRefresh = wx.Button(
            self, wx.ID_ANY, "Обновить информацию по совещаниям\n" +
            "Данные актуальны на " + str(datetime.today())[10:16],)
        self.CommonVbox.Add(self.buttonRefresh, -1, wx.EXPAND|wx.ALL, 4)
        self.buttonRefresh.Bind(wx.EVT_BUTTON, self.Refresh)
        self.buttonRefresh.SetBackgroundColour(wx.Colour(self.colour[1]))
        self.buttonRefresh.SetMinSize((225, 40))
        #print(spisoksov[5])

        # Add date menu
        #Movement buttons and others
        #Preparing Images
        MoveSize = (25, 25)
        TextSize = (125, 25)
    
        Image = ["btnarrowleft.png", "btnarrowright.png"]
        BMPs = []
       
        for i in range (0, 2):
            Im = wx.Image(os.getcwd() + "\\images\\" + Image[i]).ConvertToBitmap()
            BMPs.append(ScaleBitmap(Im, MoveSize))

        LeftBtn = wx.BitmapButton(self, wx.ID_ANY, BMPs[0])
        LeftBtn.Bind(wx.EVT_BUTTON, self.MoveLeft)

        RightBtn = wx.BitmapButton(self, wx.ID_ANY, BMPs[1])
        RightBtn.Bind(wx.EVT_BUTTON, self.MoveRight)

        self.NowDate = wx.TextCtrl(self, wx.ID_ANY, self.Date, style = wx.TE_CENTRE)
        self.NowDate.SetFont(wx.Font(14, wx.SWISS, wx.NORMAL, wx.NORMAL))
        self.NowDate.SetBackgroundColour(wx.Colour(self.colour[1]))
        self.NowDate.Bind(wx.EVT_CHAR_HOOK, self.BlockNonNumbers)
        self.NowDate.SetMinSize(TextSize)

        ButtonVbox.Add(LeftBtn, -1, wx.ALIGN_CENTRE|wx.ALL, 1)
        ButtonVbox.Add(self.NowDate, -1, wx.EXPAND|wx.ALL, 1)
        ButtonVbox.Add(RightBtn, -1, wx.ALIGN_CENTRE|wx.ALL, 1)

        self.CommonVbox.Add(ButtonVbox, -1, wx.EXPAND| wx.ALL, 0)

        # Adding meeting list
        try:
            self.CountLabels()
        except Exception as Err:
            self.Labels = self.spisoksov[5]
            ToLog("Error in CountLabels, Error code = " + str(Err))

        self.BtnId = []
        self.Btn = []
        self.RunningId = False
        i = 0        
        while i < len(self.spisoksov[5]):
            self.NameButton = wx.Button(self, wx.ID_ANY, self.Labels[i])
            self.CommonVbox.Add(self.NameButton, -1, wx.EXPAND|wx.LEFT|wx.RIGHT, 4)
            self.NameButton.Bind(wx.EVT_BUTTON, self.otclick)
            self.NameButton.Bind(wx.EVT_RIGHT_DOWN, self.RClick)
            self.NameButton.SetBackgroundColour(wx.Colour(self.colour[1]))
            self.BtnId.append(self.NameButton.GetId())
            self.Btn.append(self.NameButton)
            i = i + 1

        # Добавление кнопки лицензии
        self.buttonlic = wx.Button(self, wx.ID_ANY, "Лицензия")
        self.CommonVbox.Add(self.buttonlic, -1, wx.EXPAND|wx.ALL, 4)
        self.buttonlic.Bind(wx.EVT_BUTTON, self.license)
        self.buttonlic.SetBackgroundColour(wx.Colour(self.colour[1]))
                  
        # Добавление версии
        global MyDate
        self.version = wx.StaticText(
            self, label = "Версия от " + MyDate)
        self.CommonVbox.Add(self.version, -1, wx.ALIGN_CENTRE, 0)

        self.MonitMenu = wx.NewIdRef(count = 1)

        # Добавляем слушателя
        pub.subscribe(self.UpdateDisplay, "MainFrame")

        # Действия при нажатии кнопки закрыть
        #self.Bind(wx.EVT_CLOSE, self.OnCloseWindow)

    #=====================================================================================
    # Counting Labels
    def CountLabels(self):
        self.Labels = []
        # now I must delete "None"
        temp = []
        for i in range (0, len(self.spisoksov[6])):
            if self.spisoksov[6][i] == ["None"]:
                temp.append([])
            else:
                temp.append(self.spisoksov[6][i][:])
        self.spisoksov[6] = temp[:]
        
        # if it's Yprava
        if self.FactList != False:
            for k in range (0, len(self.spisoksov[5])):
                #print("Working with label = " + self.spisoksov[5][k])
                
                if len(self.spisoksov[5][k]) >= 41:
                    self.Labels.append(
                        self.spisoksov[5][k][:33] + "..." + self.spisoksov[5][k][-7:] + "  (" +
                        str(len(self.FactList[k])) + "/" + str(len(self.spisoksov[6][k])) + ")")
                else:
                    self.Labels.append(
                        self.spisoksov[5][k][:] + "  (" +
                        str(len(self.FactList[k])) + "/" + str(len(self.spisoksov[6][k])) + ")")
        # it's not Yprava
        else:
            for k in range (0, len(self.spisoksov[5])):
                if len(self.spisoksov[5][k] )>= 41:
                    self.Labels.append(
                        self.spisoksov[5][k][:33] + "..." + self.spisoksov[5][k][-7:])
                else:
                    self.Labels.append(self.spisoksov[5][k])
        #print("Here Labels")
        #for label in self.Labels:
        #    print(label)
            
        return 
    #=====================================================================================
    # Block NonNumbers to TEXTCTRL
    def BlockNonNumbers(self, event):
        try:
            #print("Pressed button with id = " + str(event.GetId()))
            #print("Value of this button = " + str(self.frame.FindWindowById(event.GetId()).GetValue()))
            key_code = event.GetKeyCode()
            #print("Entered = " + chr(key_code))
            #print("stringdigits = " + string.digits)
            # filter unicode characters
            if key_code == wx.WXK_NONE:
                pass
            elif key_code == wx.WXK_RETURN or key_code == 370:
                print("Paremeters entered, date = " + str(self.NowDate.GetValue()))
                self.NewDate = str(self.NowDate.GetValue())
                try:
                    self.RenewWindow(self.NewDate)
                    ToLog("Изменение даты выгрузки на " + self.NewDate)
                except Exception as Err:
                    wx.MessageBox("Failed to Renew Window, Error code = " + str(Err), "Ошибка", wx.OK)
                    ToLog("Ошибка обновления окна, код ошибки = " + str(Err))
                event.Skip()
            elif chr(key_code) in string.digits or 324 <= key_code <= 333 or key_code == 46 or key_code == 391:
                event.Skip()
            # allow special, non-printable characters
            elif (
                chr(key_code) not in string.printable and
                key_code < 300):
                print("Entered = " + str(key_code) + " or chr = " + chr(key_code))
                event.Skip()        
            else:
                print("Entered something non-number = " + str(key_code) + " or chr = " + chr(key_code))
        except Exception as Err:
            ToLog("Ошибка при обработке события в тестовой строке")
            event.Skip()   
        return
    #=======================================================================================
    def RClick(self, evt):
        #right-click on Btn
        #print("Rclicked")
        global MonitOpen
        try:
            menu = wx.Menu()
            #print("GetId = " + str(evt.GetId()))
            #print("RunningID = " + str(self.RunningId))

            menu.Append(self.MonitMenu, "Мониторинг списков присутствующих")
            self.Id = evt.GetId()
            self.Bind(wx.EVT_MENU, self.DoStartMonit, id = self.MonitMenu)
                                 
            self.clickedBtn = evt.GetEventObject().GetLabel()
            
            return self.PopupMenu(menu)
        except Exception as Err:
            ToLog("Error of Popup menu, Error code = " + str(Err))
            #print("Error of Popup menu because of = " + str(Err))
        return
    
    #============================================================================
    def DoStartMonit(self, evt):
        #print ("Pressed button = " + self.clickedBtn)
        global MonitOpen
        if self.RunningId != False or MonitOpen == True:
            wx.MessageBox("Сначала завершите мониторинг другого совещания")
            return
        for i in range (0, len(self.Labels)):
            if self.clickedBtn == self.Labels[i]:
                #print("It's meeting number = " + str(i))
                #print ("Init = ", self.spisoksov[3][i])
                #print ("id = ", self.spisoksov[2][i])
                self.Btn[i].SetBackgroundColour(wx.Colour(self.colour[2]))

                #creating monitoring frame
                self.MonitFrame = MonitFrame(
                    Label = "Анализ изменения учаcтников совещания " + self.spisoksov[5][i],
                    Id = self.spisoksov[2][i],
                    DocDir = self.DocDir)
                
                #starting thread of monitoring
                global Thread
                Thread = MonitThread(
                    [self.spisoksov[2][i], self.Date, self.spisoksov[4][i], self.spisoksov[3][i]],
                    DocDir = self.DocDir)
                Thread.setDaemon(True)
                Thread.start()

                # Функция формирования файла, исходя из выбранного совещания
                #formfile(self.spisoksov[2][i], self.Date, self.spisoksov[4][i], self.spisoksov[3][i], monitoring = True)
    #=======================================================================================================
    def SetColour(self, Id):
        try:
            for i in range (0, len(self.spisoksov[2])):
                if str(Id) == self.spisoksov[2][i]:
                    print("matched id = " + str(Id))
                    self.Btn[i].SetBackgroundColour(wx.Colour(self.colour[2]))

        except Exception as Err:
            pass
            #ToLog("Error to change colour of monit, Error code = " + str(Err))

    #============================================================================
    def DoStopMonit(self, evt):
        #print("Pressed button = " + self.clickedBtn)
        for i in range (0, len(self.Labels)):
            if self.clickedBtn == self.Labels[i]:
                #print("It's meeting number = " + str(i))
                #print ("Init = ", self.spisoksov[3][i])
                #print ("id = ", self.spisoksov[2][i])
                StopThread()
                self.Btn[i].SetBackgroundColour(wx.Colour(self.colour[1]))

                global MonitOpen
                if MonitOpen == True:
                    try:
                        wx.CallAfter(
                            pub.sendMessage, "MonitFrame",
                            mess = DestroyFrame)
                    except Exception as Err:
                        ToLog("Error closing Monitoring Frame, Error code = " + str(Err))               

    #===========================================================================
    def MoveLeft(self, evt):
        self.ChangeDate(False)
        ToLog("Изменение даты выгрузки на -1")

    def MoveRight(self, evt):
        self.ChangeDate(True)
        ToLog("Изменение даты выгрузки на +1")

    def ChangeDate(self, direction):
        try:
            DateBegin = datetime.strptime(self.Date, "%d.%m.%Y")
            if direction == True:
                DateEnd = DateBegin + timedelta(days = 1)
            else:
                DateEnd = DateBegin - timedelta(days = 1)

            NewDate = str(DateEnd)[8:10] + "." + str(DateEnd)[5:7] + "." + str(DateEnd)[0:4]
            #print("DateEnd = " + NewDate)
            self.RenewWindow(NewDate)

        except Exception as Err:
            SomeError(None, "Ошибка переключения окна на другую дату")
            ToLog("Error in ChangeDate, Error code = " + str(Err))

    #===========================================================================
    def RenewWindow(self, date):
        global Region
        try:
            #меняем название кнопки, считываем дату и положение окна
            self.buttonRefresh.SetLabel("Информация обновляется...")
            self.NewWinPos = self.frame.GetPosition()

            # считываем сетку за новую дату и создаем класс для обновления
            #self.StopThread()
            self.frame.Destroy()
            global spisoksov
            spisoksov = spisoksetka(date, str(Region))
            global WinOsn
            WinOsn = OsnWindow(None, spisoksov, WinPos = self.NewWinPos, DateUser = date, HSize = self.frame.GetClientSize()[0])
            ToLog("RenewWindow finished successfully")

        except Exception as Err:
            ToLog("Error in RenewWindow, Error code = " + str(Err))
            Error = SomeError(None, "Ошибка при обновлении окна совещаний, Error code = " + str(Err))
            #raise Exception
            
    #===========================================================================   
    # Действия при нажатой кнопке
    def otclick(self, event):
        Label = event.GetEventObject().GetLabel()
        print ("Pressed button = ", Label)
        ToLog("Нажата кнопка " + Label)
        for i in range (0, len(self.Labels)):
            if Label == self.Labels[i]:

                try:
                    if self.FactList != False:
                        if len(self.FactList[i]) < len(self.spisoksov[6][i]):
                            DifList = list(set(self.spisoksov[6][i]) - set(self.FactList[i]))
                            DifList.sort()
                            InfoDif = ""
                            for mem in DifList:
                                InfoDif = InfoDif + mem + ", "
                            InfoDif = InfoDif[:-2]
                            dlg = wx.MessageDialog(
                                self, "На данное совещание нет информации по участию руководителей:\n" +
                                InfoDif + "\n\nВсе равно выгрузить список присутствующих?",
                                "Список неполный",
                                wx.YES_NO)
                            dlg.SetYesNoLabels("&Да", "&Нет")
                            if dlg.ShowModal() != wx.ID_YES:
                                return
                                     
                    #print ("Init = ", self.spisoksov[3][i])
                    #print ("id = ", self.spisoksov[2][i])

                    # Запускаем поток выгрузки справки
                    WorkThread = ListThread(i, self.Date, self.DocDir)
                    WorkThread.start()
            
                except Exception as Err:
                    ToLog("Error in otclick, Error code = " + str(Err))
                    #raise Exception
                    
                    # запуск функции выдачи окна ошибки
                    Error = SomeError(None, "Ошибка при формировании справки, Error code = " + str(Err))
                    #raise Exception
            #else:
            #    print("Its not label = " + spisoksov[5][i])

        if Label == "OK":
            WinLicense.destroy()

        if Label == "ВЫХОД":
            self.Destroy()
            sys.exit()

        return
    #===========================================================================
        # Действия при обновлении окна
    def Refresh(self, event):
        self.RenewWindow(self.Date)
            
    #===========================================================================        
    # Лицензия
    def license(self, evt):
        ToLog ("License button pressed")
        LICENSE = (
            "Данная программа является свободным программным обеспечением\n"+
            "Вы вправе распространять её и/или модифицировать в соответствии\n"+
            "с условиями версии 2 либо по Вашему выбору с условиями более\n"+
            "поздней версии Стандартной общественной лицензии GNU, \n"+
            "опубликованной Free Software Foundation.\n\n\n"+
            "Эта программа создана в надежде, что будет Вам полезной, однако\n"+
            "на неё нет НИКАКИХ гарантий, в том числе гарантии товарного\n"+
            "состояния при продаже и пригодности для использования в\n"+
            "конкретных целях.\n"+
            "Для получения более подробной информации ознакомьтесь со \n"+
            "Стандартной Общественной Лицензией GNU.\n\n"+
            "Данная программа написана на Python\n"
            "Особая благодарность за помощь в создании Анкешеву А.Д.\n\n"+
            "Автор: Титовский С.А.")
        # Создание диалогового окна
        wx.MessageBox(LICENSE, "Лицензия", wx.OK)
    #===========================================================================
    #UpdateDisplay
    def UpdateDisplay(self, mess):
        #print("mess to MainFrame = " + str(mess))
        try:
            if mess == "I was closed":
                for btn in self.Btn:
                    btn.SetBackgroundColour(wx.Colour(self.colour[1]))
                self.RunningId = False
                self.StopThread()

            if isinstance(mess, list):
                if mess[0] == "MyId":
                    self.SetColour(mess[1])
        except Exception as Err:
            ToLog("Error in UpdateDisplay in MainFrame, Error code = " + str(mess))
                

#===============================================
#===============================================
#===============================================
#===============================================
#Frame for monitoring
class MonitFrame(wx.Frame):
    def __init__(self, Label, Id, DocDir = os.getcwd()):
        
        self.ID = Id
        wx.Frame.__init__(
            self, None, -1, Label)
        #,
        #    style = wx.MINIMIZE_BOX|wx.MAXIMIZE_BOX|wx.RESIZE_BORDER|wx.CAPTION|wx.SYSTEM_MENU|wx.CLOSE_BOX|wx.CLIP_CHILDREN)

        frameIcon = wx.Icon(os.getcwd() + "\\images\\IconPNG.png")
        self.SetIcon(frameIcon)

        global MonitOpen
        MonitOpen = True
        
        self.panel = MonitPanel(self, Id = self.ID, DocDir = DocDir)
        self.Center()
        self.Show(True)
        self.SetClientSize(400, 500)

        self.Bind(wx.EVT_CLOSE, self.OnCloseWindow)
#=========================================
    def OnCloseWindow(self, evt):
        try:
            dlg = wx.MessageDialog(
                self, "Сохранить результаты мониторинга перед закрытием?", " ", wx.YES_NO)
            dlg.SetYesNoLabels("&Да", "&Нет")
            if dlg.ShowModal() != wx.ID_YES:
                ToLog("User decided not to save results of monitoring")
            else:
                ToLog("Saving results of monitoring")
                self.panel.DoSaveLog()
        
            ToLog("MonitoringFrame Closed")
            global MonitOpen
            MonitOpen = False
            StopThread()
            evt.Skip()

        except Exception as Err:
            ToLog("Error in OnClWindow MonitFrame Error code = " + str(Err))
            evt.Skip()

# panel
class MonitPanel(wx.Panel):

    def __init__(self, parent, Id, DocDir = os.getcwd()):
        wx.Panel.__init__(self, parent = parent)
        self.ID = Id
        self.frame = parent
        self.DocDir = DocDir
        self.colour = LoadColours(self.DocDir)

        self.SetBackgroundColour(wx.Colour(self.colour[0]))
        self.CommonVbox = wx.FlexGridSizer(rows = 2, cols = 1, hgap = 6, vgap = 6)
        self.CommonVbox.AddGrowableCol(0, 1)
        self.CommonVbox.AddGrowableRow(0, 0)

        #txt for logs
        self.TextLog = wx.TextCtrl(
            self, wx.ID_ANY, " ",
            style = wx.TE_MULTILINE|wx.TE_READONLY|wx.TE_RICH)
        
        self.TextLog.SetFont(wx.Font(10, wx.SWISS, wx.NORMAL, wx.NORMAL))
        self.TextLog.SetBackgroundColour(wx.Colour(self.colour[1]))
        self.CommonVbox.Add(self.TextLog, -1, wx.EXPAND|wx.ALL, 4)
        
        #btn save
        self.BtnSave = wx.Button(self, wx.ID_ANY, "Сохранить в файл")
        self.CommonVbox.Add(self.BtnSave, -1, wx.EXPAND|wx.ALL, 4)
        self.BtnSave.Bind(wx.EVT_BUTTON, self.SaveLog)
        self.BtnSave.SetBackgroundColour(wx.Colour(self.colour[1]))
        self.BtnSave.SetMinSize((100, 40))

        self.SetSizer(self.CommonVbox)
        self.Fit()

        # Добавляем слушателя
        pub.subscribe(self.UpdateDisplay, "MonitFrame")
        self.frame.Bind(wx.EVT_CLOSE, self.OnCloseWindow)


    def SaveLog(self, evt):
        self.DoSaveLog()
        
    def DoSaveLog(self):
        try:
            info = self.TextLog.GetValue()
            DialogSave = wx.FileDialog(
                self,
                "Сохранить результаты в файл",
                defaultDir = DocDir + "\\Logs",
                wildcard = "txt files (*.txt)|*txt",
                style = wx.FD_SAVE | wx.FD_OVERWRITE_PROMPT)
            if DialogSave.ShowModal() == wx.ID_CANCEL:
                return
            else:
                try: 
                    UserPath = DialogSave.GetDirectory()
                    if DialogSave.GetFilename().find(".txt") != -1:
                        file = open(UserPath + "\\" + DialogSave.GetFilename(), "w")
                    else:
                        file = open(UserPath + "\\" + DialogSave.GetFilename() + ".txt", "w")
                except Exception as Err:
                    ToLog("Ошибка при выборе файла для сохранения лога, код ошибки = " + str(Err))
                    wx.MessageBox("Ошибка при выборе файла для сохранения лога, код ошибки = " + str(Err), " ", wx.OK)
                    return
                    
            file.write(info)
            file.close()
            
        except Exception as Err:
            self.ToLog("Ошибка при сохранении лога совещания, код ошибки = " + str(Err))
            wx.MessageBox("Ошибка при сохранении лога совещания, код ошибки = " + str(Err), " ", wx.OK)
        return

    def OnCloseWindow(self, evt):
        try:
            wx.CallAfter(pub.sendMessage, "MainFrame", mess = "I was closed")
            ToLog("Окно мониторинга было закрыто")
        except Exception as Err:
            ToLog("Ошибка при закрытии окна мониторинга, код ошибки = " + str(Err))
        evt.Skip()

    def UpdateDisplay(self, mess):
        #print("mess to MonitFrame = " + str(mess))
        try:
            self.TextLog.AppendText("\n" + str(mess))
            
        except Exception as Err:
            pass
           
#=============================================
#=============================================
#=============================================
#=============================================
# Monitoring thread 
class MonitThread(threading.Thread):
    #def __init__(self, parent, spisoksov):
    def __init__(self, spisok, DocDir = os.getcwd()):
        super().__init__()
        #import parameters
        #self.parent = parent
        self.Id = spisok[0]
        self.Date = spisok[1]
        self.Time = spisok[2]
        self.Init = spisok[3]
        self.DocDir = DocDir
        
        global TimeRefresh
        self.TimeSleep = TimeRefresh
        self.Stop = False

        sound = SoundThread(self.DocDir + "\\Sounds\\sound.wav")
        sound.setDaemon(True)
        sound.start()

#============================================================================
    def run(self):
        try:
            print("Downloading current data")
            
            self.ToFrame(
                "Начало мониторинга списков совещания " + self.Init +
                " в " + self.Time + " за дату " + self.Date + ", время обновления = " + str(self.TimeSleep) + " секунд")
            ToLog(
                "Начало мониторинга совещания " + self.Init +
                " в " + self.Time + " за дату " + self.Date + ", время обновления = " + str(self.TimeSleep) + " секунд")
            self.ToFrame("\tУчастники отображаются в формате:", ShTime = False)
            self.ToFrame("\tДолжность | Участник (Ф.И.О.) | Примечание | Студия", ShTime = False)
            self.FirstList = formfile(self.Id, self.Date, self.Time, self.Init, monitoring = True)
            
            #print("Here's FirstList")
            #for data in self.FirstList:
            #    print(str(data))
        except Exception as Err:
            self.Stop = True
            ToLog(
                "Ошибка при выгрузке начального списка при запуске мониторинга совещания " + self.Init +
                " в " + self.Time + " за дату " + self.Date + ", код ошибки = " + str(Err))
            #raise Exception

        while self.Stop == False:
            for i in range (0, self.TimeSleep):
                time.sleep(1)
                wx.CallAfter(
                    pub.sendMessage, "MainFrame",
                    mess = ["MyId", self.Id])
                if self.Stop == True:
                    break
                # end of cycle
                print("iteration of sleeptime = " + str(i))
                if i == self.TimeSleep - 1:
                    self.CompareLists()

        ToLog("End of monitoring Thread")
#============================================================================
    def CompareLists(self):
        try:
            self.NewList = formfile(self.Id, self.Date, self.Time, self.Init, monitoring = True)
            #print("Here's FirstList")
            #for data in self.FirstList:
            #    print(str(data))

            #print("Here's NewList")
            #for data in self.NewList:
            #    print(str(data))

            if self.FirstList == self.NewList:
                print("Списки не изменились")
                #self.ToFrame("Списки не изменились")
                
            else:
                ToLog("Зафиксированы изменения в списках")
                sound = SoundThread(self.DocDir + "\\Sounds\\sound.wav")
                sound.setDaemon(True)
                sound.start()
                

                ToLog("First = " + str(self.FirstList))
                ToLog("New = " + str(self.NewList))
                
                #self.ToFrame("Появились изменения в списках")
                #block of comparing dif lists
                changedReg = []
                
                for regi in range (0, len(self.FirstList)):
                    if self.FirstList[regi] == self.NewList[regi]:
                        pass
                    else:
                        changedReg.append(regi)
                #nameReg = []
                stringReg = ""
                print("")
                #for all changed regions
                for regi in range (0, len(changedReg)):
                    if changedReg[regi] == 0:
                        #nameReg = "Управлении"
                        stringReg = "Управления "
                    elif changedReg[regi] == 1:
                        #nameReg = "Омском территориальном управлении"
                        stringReg = "Омского территориального управления "
                    elif changedReg[regi] == 2:
                        #nameReg = "Новосибирском территориальном управлении"
                        stringReg = "Новосибирского территориального управления "
                    elif changedReg[regi] == 3:
                        #nameReg = "Кузбасском территориальном управлении"
                        stringReg = "Кузбасского территориального управления "
                    elif changedReg[regi] == 4:
                        #nameReg = "Алтайском территориальном управлении"
                        stringReg = "Алтайского территориального управления "

                    #sorting raw list to format [[studia1, studia2],[[dolg1],[dolg2]], [[members1], [members2]]]
                    temp = [[],[],[]]
                    List1 = self.FormattingLists(self.FirstList[changedReg[regi]])[:]
                    List2 = self.FormattingLists(self.NewList[changedReg[regi]])[:]

                    if set(List1) != set(List2):
                    #if List1 != List2:
                        self.ToFrame("Появились изменения в списках " + stringReg)

                    #format from raw list [[studia + dolgn], [fio], [prim]]
                    #to format [[studia1, studia2],[[dolg1],[dolg2]], [[members1], [members2]]]

                    # if no stud changed
                    #if set(List1[0]) == set(List2[0]):
                        #self.ToFrame("\tИзменений в составе студий нет")

                    #now comparing dolg
                    addedMem = list(set(List2) - set(List1))
                    if isinstance(addedMem, list):
                        addedMem.sort()
                        for mem in addedMem:                      
                            self.ToFrame("\tДобавлен участник " + mem, ShTime = False)
                    removedMem = list(set(List1) - set(List2))
                    if isinstance(removedMem, list):
                        removedMem.sort()
                        for mem in removedMem:
                            self.ToFrame("\tУдален участник " + mem, ShTime = False)

                self.FirstList = self.NewList
                                                             
        except Exception as Err:
            self.Stop = True
            ToLog(
                "Ошибка при сравнении списков при мониторинге совещания " + self.Init +
                " в " + self.Time + " за дату " + self.Date + ", код ошибки = " + str(Err))
        
            #raise Exception

#===============================================================================
    #format from raw list [[studia + dolgn], [fio], [prim]]
    #to format [[studia1, studia2],[[dolg1],[dolg2]], [[members1], [members2]]]
    #to format[[dolg|fio|prim|stud],[],...]
    def FormattingLists(self, List):
    
        #print("List before func = " + (str(List)))
        temp = []
        for i in range (0, len(List[0])):
            #finding studia
            if List[0][i].find("КАБИНЕТ") != -1:
                #if i == len(List[0]):
                #    break
                #temp[0].append(List[0][i][7:])  
                for k in range (i + 1, len(List[0])):
                    if List[0][k].find("КАБИНЕТ") != -1:
                        break
                    #tempDolg.append(List[0][k])
                    temp.append(List[0][k] + " | " + List[1][k] + " | " + List[2][k] + " | из студии " + List[0][i][7:])
                    
        #print("List after func = " + (str(temp)))
        return temp
                
#===============================================================================
    def ToFrame(self, mess, ShTime = True):
        if ShTime == True:
            mess = "\n" + str(datetime.today())[10:19] + "  " + str(mess)
        try:
            wx.CallAfter(
                pub.sendMessage, "MonitFrame",
                mess = mess)
        except Exception as Err:
            ToLog("Ошибка отправки сообщения в окно мониторинга, код ошибки = " + str(Err))
            #raise Exception
        
#===============================================
#===============================================
#===============================================
#===============================================        
# Создание класса окна любой ошибки
def SomeError(parent, title):
    wx.MessageBox(title, "Ошибка", wx.OK)

#===============================================
#===============================================
#===============================================
#===============================================
# Составление списка совещаний        
def spisoksetka(date, region):
  
    # Подставляем дату в ссылку:
    ssilka = str(
        "http://10.132.71.156/pls/ss/selector.sels.list?us=" + 
        str(region) +
        "&str=" + 
        date +
        "&wday=5")

    # Запрашиваем код
    try:
        response = requests.get(ssilka)
    except Exception as Err:                   
        ToLog("Ссылка селекторных для данной даты недоступна, код ошибки = " + str(Err))
        sys.exit()

    idsov = []
    namesov = []
    rezhimsov = []
    timesov = []
    initsov = []
    spisoksov = []
    spisokuchast = []
    nomsov = 1

    # Если Управление
    if int(region) == 0:
        
        filesplit = response.text.splitlines()

        #Задаем параметры поиска
        poisk = "&us=0&sid="
        poisk2 = '''<td width=15% class=zag>Примечание</td>'''
        poisk3 = '''&nbsp;</td></tr>'''
        poisk4 = '''<td class="zag" rowspan=2>'''
        poisk5 = '''<td class="msk" rowspan=2>'''
        poisk6 = '''<a href="javascript:go(0,1,0'''
        poisk7 = '''Регион-'''
        poisk8 = '''&nbsp;'''
        nachalo = "2"
        konec = '''</td>'''

        #Обработка кода страницы и составление списков
        for i in range(0, len(filesplit)-1):

            filesplit[i] = str(filesplit[i]).strip()

            #добавление в списки разделителей - строки Начало совещания и Список участников
            if (
                filesplit[i].find('''<td width=15% class=zag>Примечание</td>''')>-1
                or
                filesplit[i]=='''&nbsp;</td></tr>'''):
                spisoksov.append(str(nomsov))
                spisokuchast.append("Список участников  "+str(nomsov))
                nomsov = nomsov + 1

            #составление списка участников конференций (необработанного)
            if filesplit[i].find('''<a href="javascript:go(0,1,0''')!=-1:
                spisokuchast.append(filesplit[i][filesplit[i].find('''">''')+2:filesplit[i].find('''</a>''')])
 
            #составление списка как в SMS
            if (
                (filesplit[i].find(poisk4)!=-1)
                or
                (filesplit[i].find(poisk5)!=-1)):
                if filesplit[i].find('''<br>''') > -1:
                    filesplit[i] = filesplit[i][:filesplit[i].find('''</td>''')+1]
                filesplit[i] = filesplit[i][filesplit[i].find(nachalo) + 2:filesplit[i].find(konec)]
                spisoksov.append(filesplit[i])

            #составление списка ID конференций (внутри списка SMS)
            if filesplit[i].find(poisk)!=-1:
                filesplit[i] = filesplit[i][filesplit[i].find(poisk)+10:filesplit[i].find('''>"''')-1]
                spisoksov.append(str(filesplit[i]))

        for i in range (5, len(spisoksov)):
            if (i+1)%6==0:
                rezhimsov.append(spisoksov[i-3])
                timesov.append(spisoksov[i-2])
                initsov.append(spisoksov[i-1])
                idsov.append(spisoksov[i])
                namesov.append(str(str(spisoksov[i-1])+" в "+str(spisoksov[i-2][:spisoksov[i-2].find("<br")])))


    # Если не Управление (код не выверен!)
    else:
        cod = response.text.splitlines()
        
        #Обработка кода страницы и составление списков
        for i in range(0, len(cod)-1):

            cod[i] = str(cod[i]).strip()
            #print(cod[i])

            #составление списка как в SMS

            if (
                cod[i].find('''<td class="msk1" rowspan=2>''')!=-1
                or
                cod[i].find('''<td class="msk" rowspan=2>''')!=-1
                or
                cod[i].find('''<td class="zag" rowspan=2>''')!=-1):

                if cod[i].find('''<br>''')!=-1:
                    cod[i] = cod[i][cod[i].find('''rowspan=2''')+10:cod[i].find('''<br>''')]
                else:
                    cod[i] = cod[i][cod[i].find('''rowspan=2>''')+10:cod[i].find('''</td>''')]
                #print("Кусочек совещания = " + cod[i])
                spisoksov.append(cod[i])

            if (
                cod[i].find('''&us='''+str(region))!=-1
                and
                cod[i].find('''Справка о присутствующих''')!=-1):

                cod[i] = cod[i][cod[i].find('''id=''')+3:cod[i].find('''&us=''' + str(region))]
            
                #print("Нашел id = " + cod[i])
                idsov.append(cod[i])

        for i in range (3, len(spisoksov), 4):
            namesov.append(str(str(spisoksov[i])+" в "+str(spisoksov[i-1])))
            #print("Added sov = " + str(str(spisoksov[i])+" в "+str(spisoksov[i-1])))
            initsov.append(spisoksov[i])
            rezhimsov.append(spisoksov[i-2])
            timesov.append(spisoksov[i-1])
                
        #for i in range (0, len(idsov)):
           # print (str(i+1) + ".  " + namesov[i] + " его id = " + idsov[i])


    # формируем списки с учетом отмен и проверок
    idsov1 = []
    namesov1 = []
    rezhimsov1 = []
    initsov1 = []
    timesov1 = []
    uchastsov1 = []
    nomer = []
    nomernach = 1
    uchastsov = []
    temp_uchast = []

    # Формируем список списков участников

    for i in range (1, len(spisokuchast)):
        if spisokuchast[i].find("Список участников")==-1:
            temp_uchast.append(spisokuchast[i])
        else:
            if len(temp_uchast)==0:
                uchastsov.append(["None"])
                temp_uchast.clear()
            else:
                uchastsov.append(temp_uchast[:])
                temp_uchast.clear()
                    

    for i in range (0, len(namesov)):
        if (
            rezhimsov[i].find('''тмена''')!=-1
            or
            rezhimsov[i].find('''роверка''')!=-1):
            continue
        else:
            nomernach = nomernach + 1
            idsov1.append(idsov[i])
            rezhimsov1.append(rezhimsov[i])
            namesov1.append(namesov[i])
            if len(timesov[i]) < 16:
                timesov[i] = timesov[i].replace("<br>-<br>", "")
                if timesov[i][-1] == ":":
                    timesov[i] = timesov[i][:-1]
                timesov1.append(timesov[i])
                             
            else:
                timesov1.append(timesov[i].replace("<br>-<br>", "-"))
            initsov1.append(initsov[i])
            nomer.append(str(nomernach-1))
            if Region == 0:
                uchastsov1.append(uchastsov[i])
            
    idsov = idsov1
    namesov = namesov1
    rezhimsov = rezhimsov1
    timesov = timesov1
    initsov = initsov1
    uchastsov = uchastsov1
    labelkn = []

    # Выводим список совещаний на дату
    for i in range(0, len(namesov)):
        #print (str(i+1)+".  "+namesov[i] + " список участников = " + str(uchastsov[i]))
        labelkn.append(str(i+1)+".  "+str(namesov[i]))

    #for Time in timesov:
    #    Time = Time.replace("<br>-<br>", "")

    itog = []
    itog.append(nomer)
    itog.append(namesov)
    itog.append(idsov)
    itog.append(initsov)
    itog.append(timesov)
    itog.append(labelkn)
    itog.append(uchastsov)

    return itog

#===============================================
#===============================================
#===============================================
#=============================================== 
# Функция формирования файла, исходя из выбранного совещания
def formfile(idsov, datass, timess, initss, monitoring = False, DocDir = os.getcwd(), Dlg = None):
    #print("Params")
    #print(idsov)
    #print(datass)
    #print(timess)
    #print(initss)
    
    itogitogov = []
    for k in range (0,5):
        ssilkaSS = str(
            "http://10.132.71.156/pls/ss/selector.report.study_p?sid="+idsov+"&us="+str(k))
        if monitoring == False:
            # отправка обновения основному окну
            if k == 0:
                Dlg.Update(Dlg.GetValue() + 1, "Запрашиваю данные по Управлению...")
            elif k == 1:
                Dlg.Update(Dlg.GetValue() + 1, "Запрашиваю данные по Омскому региону...")
            elif k == 2:
                Dlg.Update(Dlg.GetValue() + 1, "Запрашиваю данные по Новосибирскому региону...")
            elif k == 3:
                Dlg.Update(Dlg.GetValue() + 1, "Запрашиваю данные по Кузбасскому региону...")
            elif k == 4:
                Dlg.Update(Dlg.GetValue() + 1, "Запрашиваю данные по Алтайскому региону...")
            time.sleep(0.3)
                
        # Запрашиваем код
        try:
            responseSS = requests.get(ssilkaSS)
        except Exception as Err:
            ToLog("Ошибка в обработке страницы списка присутствующих, код ошибки = " + str(Err))
            sys.exit()
        filesplit = responseSS.text.splitlines()

        #Задаем наши списки и номер совещания
        zagolovok = []
        dolgnost = []
        fio = []
        prim = []
        if k==0:
            zagolovok.append("Справка")
            zagolovok.append("о присутствующих на селекторном совещании")
            zagolovok.append("Дата проведения:  "+str(datass)+"                                       Время:  "+str(timess))
            zagolovok.append(" ")
            zagolovok.append("Совещание назначено: "+str(initss))

        if monitoring == False:
            Dlg.Update(Dlg.GetValue() + 1, "Обрабатываю полученные данные")
            time.sleep(0.3)

        #Обработка кода страницы и составление списков
        for i in range(0, len(filesplit)-1):

            # Формируем тему совещания
            if (
                filesplit[i].find('''<td colspan=2 class=m1 >Тема:''')!=-1
                and
                k==0):
                zagolovok[3] = "Тема: "
                if filesplit[i].find('''</td>''')!=-1:
                    zagolovok[3] = zagolovok[3] + str(filesplit[i][filesplit[i].find('''&nbsp;&nbsp''')+12:filesplit[i].find('''</td>''')])
                else:
                    n = i
                    while filesplit[n].find('''</td>''')==-1:
                        if filesplit[n].find('''<td colspan=2 class=m1 >Тема:''')!=-1:
                            zagolovok[3] = zagolovok[3] + filesplit[n][filesplit[i].find('''&nbsp;&nbsp''')+12:]
                        else:
                            zagolovok[3] = zagolovok[3] + filesplit[n]
                        n = n+1
                    zagolovok[3] = zagolovok[3]+str(filesplit[n][:filesplit[n].find('''</td>''')])

                # Замена служебных символов (кривые руки)
                zagolovok[3] = zagolovok[3].replace('''\x02''', "-")
                                    
                # Формируем списки для таблицы - должность, ФИО, Примечание
            if (
                filesplit[i].find('''<tr><td colspan=3 class=z2>''')!=-1
                or
                filesplit[i].find('''<tr><td class=spr valign=top>''')!=-1):
                
                if filesplit[i].find('''<tr><td colspan=3 class=z2>''')!=-1:
                    dolgnost.append("КАБИНЕТ" + str(filesplit[i][filesplit[i].find('''z2>''')+3:filesplit[i].find('''</td>''')]))
                    fio.append("NONE")
                    prim.append("NONE")

                if filesplit[i].find('''<tr><td class=spr valign=top>''')!=-1:
                    dolgnost.append(filesplit[i+1])
                    fio.append("Новые участники:")
                    n = i
                    while filesplit[n].find('''</table></td>''')==-1:
                        n = n+1
                    for s in range (i,n):
                       
                        if filesplit[s].find('''<td class=spr>''')!=-1:
                            fio.append(
                                str(filesplit[s])[filesplit[s].find('''<td class=spr>''')+14:filesplit[s].find('''&nbsp;&nbsp''')]+
                                str("  ")+str(filesplit[s])[filesplit[s].find('''&nbsp;&nbsp''')+12:filesplit[s].find('''</td>''')])
                        elif (
                            filesplit[s+1].find('''</table></td>''')!=-1
                            and
                            filesplit[s].find('''<table width=''')!=-1):
                            fio.append("PUSTO")
                                                
                    prim.append(filesplit[n+2])

        # Преобразовываем список участников, чтобы сгруппировать их по должностям    
        fio.append(str("Новые участники"))
        fio1 = []

        for i in range (0, len(fio)):
            if fio[i] =="NONE":
                fio1.append("NONE")
            elif (fio[i].find("Новые участники")!=-1) and (i<(len(fio)-2)):
                fio1.append(" ")
                n = i+1
                while fio[n].find("Новые участники")==-1:
                    n = n+1
                for s in range (i,n):
                    if fio[s].find("Новые участники")==-1:
                        temp = fio1[-1][:]
                        fio1[-1] = temp+"/NEXT/"+fio[s][:]

        #for i in range(0, len(prim)):
        #    print ("Долж = "+dolgnost[i]+" --- ФИО = "+fio1[i]+" --- Прим = "+prim[i])

        # итог цикла
        itog = [zagolovok, dolgnost, fio1, prim]
        itogitogov.append(itog)


    # if it's from Monitoring
    if monitoring == True:
            
        itoglist = []
        for k in range (0, 5):
            temp = [[], [], []]

            dolg = itogitogov[k][1]
            fio = itogitogov[k][2]
            prim = itogitogov[k][3]

            #for i in range (0, len(dolg)):
            #    if dolg[i].find("КАБИНЕТ") == -1:
            #        temp[0].append(dolg[i])
            #        temp[1].append(fio[i])
            #        temp[2].append(prim[i])

            #dolg = temp[0]
            #fio = temp[1]
            #prim = temp[2]

            #print("Table before format")
            #for i in range (0, len(dolg)):
            #    print(str(i) + ".  " + dolg[i] + " | " + fio[i] + " | " + prim[i])

            #for i in range (0, len(fio)):
            #    fio[i] = fio[i].replace("/NEXT/NONE", " ")
            #    fio[i] = fio[i].replace("/NEXT/PUSTO", "Нет информации")

            for i in range (1, len(dolg)+1):
                dolg[i-1] = dolg[i-1].replace("&nbsp;", " ").strip()
                fio[i-1] = fio[i-1][7:].replace("/NEXT/"," ")
                fio[i-1] = fio[i-1].replace("PUSTO", " ")
                fio[i-1] = fio[i-1].replace("NONE", "").strip()
                prim[i-1] = prim[i-1].replace("&nbsp;"," ").strip()

            #print("Table after format")
            #for i in range (0, len(dolg)):
            #    print(str(i) + ".  " + dolg[i] + " | " + fio[i] + " | " + prim[i])
            itoglist.append([dolg[:], fio[:], prim[:]])
        temp = itoglist
        return itoglist
                    
    #Not minitoring
    # Создание директории для справки
    dir_output = DocDir + "\\Conference_Lists"
    try:
        filelist = os.listdir(dir_output)
    except Exception as Err:
        wx.MessageBox("Ошибка работы с каталогом для сохранения списков, Error code = " + str(Err), " ", wx.OK)
        #raise Exception
        return

    while len(filelist) >= 8:
        if len(os.listdir(dir_output)) < 8:
            break
        try:
            os.remove(os.path.abspath(FindOldest(dir_output)))
        except Exception as Err:
            ToLog("Старый файл справки не был удален, код ошибки = " + str(Err))
            break

    # создаем наш документ
    doc = Document()

    # Задаем стиль заголовка
    obj_styles = doc.styles
    obj_charsst = obj_styles.add_style("Zagolovok", WD_STYLE_TYPE.PARAGRAPH)
    obj_font = obj_charsst.font
    obj_font.name = "Times New Roman"
    obj_font.size = Pt(14)
    obj_font.bold = True
    p = doc.styles["Zagolovok"].paragraph_format
    p.space_after = Pt(4)
    p.space_before = Pt(0)
    p.line_spacing = Pt(0)

    # Задаем студий, указанных в таблице
    obj_styles = doc.styles
    obj_charsst = obj_styles.add_style("Studia", WD_STYLE_TYPE.PARAGRAPH)
    obj_font = obj_charsst.font
    obj_font.name = "Times New Roman"
    obj_font.size = Pt(12)
    obj_font.bold = True
    p = doc.styles["Studia"].paragraph_format
    p.space_after = Pt(0)
    p.space_before = Pt(0)
    p.line_spacing = Pt(0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
    # Задаем стиль таблицы
    obj_styles = doc.styles
    obj_charsst = obj_styles.add_style("Table", WD_STYLE_TYPE.TABLE)
    obj_font = obj_charsst.font
    obj_font.name = "Times New Roman"
    obj_font.size = Pt(12)
    obj_font.bold = False
    p = doc.styles["Table"].paragraph_format
    p.space_after = Pt(2)
    p.space_before = Pt(2)
    p.line_spacing = Pt(0)
    p.left_indent = Inches(0.05)

    # Задаем стиль "итого"
    obj_styles = doc.styles
    obj_charsst = obj_styles.add_style("Itogo", WD_STYLE_TYPE.PARAGRAPH)
    obj_font = obj_charsst.font
    obj_font.name = "Times New Roman"
    obj_font.size = Pt(12)
    obj_font.bold = False

    #составление заголовка
    zagolovok = itogitogov[0][0]
    for i in range (0, len(zagolovok)):
        p = doc.add_paragraph (str(zagolovok[i]), style = "Zagolovok").alignment = WD_ALIGN_PARAGRAPH.CENTER

    #подсчет итого
        AllItogo = 0
    for k in range (0,5):
        fio = itogitogov[k][2]
        itogo = ""
        for i in range (0, len(fio)):
            fio[i]=fio[i].replace("/NEXT/NONE", " ")

            if (
                fio[i].find("/NEXT/PUSTO")==-1
                and
                fio[i]!="/NEXT/"):
                itogo = itogo+str("/Pause/")+str(fio[i])
                         
        AllItogo = AllItogo + itogo.count("NEXT")
        print("Add to ALLITOGO = " + str(itogo.count("NEXT")))

    chelovek = " человек"
    if (
        AllItogo%10==2 or AllItogo==2
        or
        AllItogo%10==3 or AllItogo==3
        or
        AllItogo%10==4 or AllItogo==4):
        chelovek = " человека"

    if (
        AllItogo==12 or AllItogo%100==12
        or
        AllItogo==13 or AllItogo%100==13
        or
        AllItogo==14 or AllItogo%100==14):
        chelovek = " человек"
                
    p = doc.add_paragraph ("\nВсего к совещанию подключено  "+str(AllItogo) + chelovek, style = "Zagolovok")
    
    #начинаем формировать таблицы в цикле
    for k in range (0, 5):
        if k == 0:
            Dlg.Update(Dlg.GetValue() + 1, "Формирую справку по Управлению...")                 
        elif k == 1:
            Dlg.Update(Dlg.GetValue() + 1, "Формирую справку по по Омскому региону...")
        elif k == 2:
            Dlg.Update(Dlg.GetValue() + 1, "Формирую справку по Новосибирскому региону...")
        elif k == 3:
            Dlg.Update(Dlg.GetValue() + 1, "Формирую справку по Кузбасскому региону...")
        elif k == 4:
            Dlg.Update(Dlg.GetValue() + 1, "Формирую справку по Алтайскому региону...")

        time.sleep(0.3)
              
        zagolovok = itogitogov[k][0]
        dolg = itogitogov[k][1]
        fio = itogitogov[k][2]
        prim = itogitogov[k][3]

        # Подсчет количества участников (пока с учетом наличия хоть какой-то информации от региона)
        itogo = ""
        for i in range (0, len(fio)):
            itogo = itogo+str("/Pause/")+str(fio[i])
                
        #print("ITOGO do = "+itogo)
        itogo = itogo.count("NEXT")
            
        if itogo>0:
            if k==1:
                p = doc.add_paragraph ("Омский регион", style = "Zagolovok").alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif k==2:
                p = doc.add_paragraph ("Новосибирский регион", style = "Zagolovok").alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif k==3:
                p = doc.add_paragraph ("Кузбасский регион", style = "Zagolovok").alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif k==4:
                p = doc.add_paragraph ("Алтайский регион", style = "Zagolovok").alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Подсчет количества участников (с учетом пустых строк)
            itogo = ""
                        
            for i in range (0, len(fio)):
                fio[i]=fio[i].replace("/NEXT/NONE", " ")

                if (
                    fio[i].find("/NEXT/PUSTO")==-1
                    and
                    fio[i]!="/NEXT/"):
                    itogo = itogo+str("/Pause/")+str(fio[i])
            
            itogo = itogo.count("NEXT")
            
            # Вставляем таблицу
            tableosn = doc.add_table(rows = len(dolg)+1, cols = 3)
            tableosn.style = "Table"

            # Выставляем ширину столбцов
            col1 = tableosn.columns[0]
            #col1.width = Cm(2.75)
            col1.width = Cm(3)
            for cell in tableosn.columns[0].cells:
                #cell.width = Cm(2.75)
                cell.width = Cm(3)
            col2 = tableosn.columns[1]
            col2.width = Cm(8.5)
            for cell in tableosn.columns[1].cells:
                cell.width = Cm(8.5)
            col3 = tableosn.columns[2]
            col3.width = Cm(4)
            for cell in tableosn.columns[2].cells:
                cell.width = Cm(4)

            tableosn.autofit = False

            # Заполняем таблицу данными
            tableosn.cell (0,0).text = "Служба\nотдел\nпредприятие"
            tableosn.cell (0,1).text = "Присутствует - Должность - ФИО"
            tableosn.cell (0,2).text = "Примечание"
            for i in range (0,3):
                tableosn.cell (0,i).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    
            for i in range (1, len(dolg)+1):
                if dolg[i-1].find("КАБИНЕТ")!=-1:
                    if dolg[i-1].find(" т.")!=-1:
                        dolg[i-1] = dolg[i-1][:dolg[i-1].find(" т.")]
                    tableosn.cell(i,0).merge(tableosn.cell(i,2))
                    if (
                        dolg[i-1].find ("Кабинет  т")!=-1
                        or
                        dolg[i-1].find ("Каналы связи совещаний")!=-1
                        or
                        dolg[i-1].find ("Аудио")!=-1
                        or
                        dolg[i-1].find("Студия Кабинет")!=-1):
                        tableosn.cell (i,0).text = "Кабинеты руководителей"

                    else:
                        tableosn.cell (i,0).text = dolg[i-1][7:]
                    for row in itertools.islice(tableosn.rows, i, i+1):
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                paragraph.style = "Studia"

                else:
                    dolg[i-1] = dolg[i-1].replace("&nbsp;", " ")
                    tableosn.cell (i,0).text = dolg[i-1].strip()
                    fio[i-1] = fio[i-1][7:].replace("/NEXT/","\n ")
                    fio[i-1] = fio[i-1].replace("PUSTO", " ")
                    fio[i-1] = fio[i-1].replace("NONE", "")
                    #tableosn.cell (i,1).text = " "+fio[i-1]
                    tableosn.cell (i,1).text = fio[i-1].strip()
                    prim[i-1] = prim[i-1].replace("&nbsp;"," ")
                    tableosn.cell (i,2).text = prim[i-1].strip()
            
            # Рисуем границы
            tableModify(tableosn)

            # Выравниваем по центру первую строку
            for row in itertools.islice(tableosn.rows, 0, 1):
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

            
            # Добавляем Итого:
            chelovek = " человек"
            if (
                itogo%10==2 or itogo==2
                or
                itogo%10==3 or itogo==3
                or
                itogo%10==4 or itogo==4):
                chelovek = " человека"

            if (
                itogo==12 or itogo%100==12
                or
                itogo==13 or itogo%100==13
                or
                itogo==14 or itogo%100==14):
                chelovek = " человек"
                
            p = doc.add_paragraph ("Итого:       "+str(itogo)+chelovek, style = "Itogo")

        #print("Table after format")
        #for i in range (0, len(dolg)):
        #    print(str(i) + ".  " + dolg[i] + " | " + fio[i] + " | " + prim[i])

    # Сохраняем файл
    nameDoc = str(initss).strip() + " в " + str(timess)
    #nameDoc = nameDoc.replace("\t", " ")
    nameDoc = nameDoc.replace(":", "-")
    
    ListFiles = os.listdir(dir_output)
    #print(nameDoc)

    try:
        num = 1
        NewName = nameDoc
        while True:
            if NewName + ".docx" in ListFiles:
                NewName = nameDoc + "(" + str(num) + ")"
                num = num + 1
                #print("File " + NewName + " in ListDir")
            else:
                #print("No such file in dir")
                break
        nameDoc = NewName + ".docx"
        #print ("\n\tСохраняю справку в файл под названием: " + nameDoc)
        doc.save(dir_output + "\\" + nameDoc)
        time.sleep(1)
        #print ("\n\tСправка успешно сохранена!!!\n\n")
        ToLog("Успешное сохранение файла со списками " + dir_output + "\\" + nameDoc)

        # ОТкрываем каталог перед пользователем
        try:
            path = os.path.realpath(dir_output)
            os.startfile(path)                        
        except Exception as Err:
            SomeError(None, "Ошибка при попытке открыть директорию...\n")
            ToLog("Ошибка при попытке открыть директорию, код ошибки = " + str(Err))
        try:
            path = os.path.realpath(dir_output + "\\" + nameDoc)
            os.startfile(path)                        
        except Exception as Err:
            SomeError(None, "Ошибка при попытке открыть файл...\n")
            ToLog("Ошибка при попытке открыть файл = " + str(Err))
    except Exception as Err:
        SomeError(
            None,
            "Ошибка в сохранении файла списка присутствующих\n" +
            "Код ошибки = " + str(Err))
        
        #raise Exception

    return

'''--------------------------------------------------------------------------------------------'''   
# Функция обработки границ таблицы
def tableModify(table):
    tbl = table._tbl
    for cell in tbl.iter_tcs():
        tcPr = cell.tcPr
        tcBorders = OxmlElement("w:tcBorders")
        top = OxmlElement("w:top")
        top.set(qn("w:val"), "thick")
        top.set(qn("w:sz"), "10")

        left = OxmlElement("w:left")
        left.set(qn("w:val"), "thick")
        left.set(qn("w:sz"), "10")
        
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "thick")
        bottom.set(qn("w:sz"), "10")
        bottom.set(qn("w:space"), "0")
        bottom.set(qn("w:color"), "auto")

        right = OxmlElement("w:right")
        right.set(qn("w:val"), "thick")
        right.set(qn("w:sz"), "10")

        tcBorders.append(top)
        tcBorders.append(left)
        tcBorders.append(bottom)
        tcBorders.append(right)
        tcPr.append(tcBorders)

    return

#=============================================
#=============================================
#=============================================
#=============================================
# Функция проверки пользователя
def AuthDlg(DocDir, Paths):
    try:
        global Region
        thr = AuthThread(DocDir, Paths)
        thr.setDaemon(True)
        thr.start()
        thr.join()

        #print("REGION = " + str(Region))
        if Region != 999:
            OpenMainFrame(DocDir)

    except Exception as Err:
        ToLog("Error in AthDlg, Error code = " + str(Err))
        #raise Exception

def OpenMainFrame(DocDir):
    try:
        global Region
        # Составляем список совещаний
        global spisoksov
        todaydate = (
            str(datetime.today())[8:10] + "." + str(datetime.today())[5:7] + "." +
            str(datetime.today())[0:4])

        ToLog("region = " + str(Region))
        spisoksov = spisoksetka(todaydate, Region)
        ToLog("Составление перечня совещаний на " + todaydate)
        ToLog("spisoksetke = " + str(spisoksov))

        global WinOsn
        WinOsn = OsnWindow(None, spisoksov, DateUser = todaydate, DocDir = DocDir)
        
    except Exception as Err:
        ToLog("Error in OpenMainFrame, Error code = " + str(Err))
        #raise Exception
    
#=============================================
#=============================================
#=============================================
#=============================================
# Thread for authorization
class AuthThread(threading.Thread):
    def __init__(self, docdir, paths, authdlg = None):
        super().__init__()
        self.stop = False
        self.DocDir = docdir
        self.paths = paths
        self.dlg = authdlg
        self.tempName = docdir + "\\Temp\\temp.db"

    def ExceptDecorator(func):
        def wrapper(self):
            ToLog("Start of " + func.__name__)
            try:
                func(self)
            except Exception as Err:
                ToLog("Error in " + func.__name__ + ", Error code = " + str(Err))
                #return False
                raise Exception
            else:
                ToLog(func.__name__ + " finished succesfully")
        return wrapper

    def run(self):
        ToLog("AuthThread started!!!")
    
        self.dlg = wx.GenericProgressDialog(
            "Авторизация",
            "Подождите, пока происходит авторизация пользователя " + str(socket.gethostname()),
            maximum = 4,
            parent = None,
            style = wx.PD_AUTO_HIDE|wx.PD_SMOOTH)
            #style = wx.PD_AUTO_HIDE|wx.PD_APP_MODAL|wx.PD_SMOOTH)
        
        self.dlg.Update(0, "Начинаю подключение к АС Селекторных...")
        time.sleep (0.5)
    
        if self.FindMyName() != False:
            self.OpenPaths()
        ToLog("AuthThread finished!!!")

    @ExceptDecorator
    def FindMyName(self):
        self.MyName = str(socket.gethostname())

    @ExceptDecorator
    def OpenPaths(self):
        global Region
        Region = int(999)
        
        #FTP
        ToLog("Try FTP")
        self.dlg.Update(1, "Подключение к АС Селекторных (попытка 1 из 3)...")
        time.sleep(1)

        self.DownLoadFTP()
        print("tempfile = " + str(self.tempFile))
        if self.tempFile == True:
            result = self.CheckSQL()
            result = self.CheckResult(result, labelnum = "1")
            if result != False:
                self.AuthSucceed()
            else:
                self.AuthFailed()

        #OPS
        else:
            ToLog("Try OPS")
            self.dlg.Update(2, "Подключение к АС Селекторных (попытка 2 из 3)...")
            time.sleep(1)
            self.CopyFromPath()
            if self.tempFile == True:
                result = self.CheckSQL()
                result = self.CheckResult(result, labelnum = "2")
                if result != False:
                    self.AuthSucceed()
                else:
                    self.AuthFailed()
            else:
                #LocalBase
                ToLog("Try LocBase")
                self.dlg.Update(3, "Подключение к АС Селекторных (попытка 3 из 3)...")
                time.sleep(1)
                result = self.CheckSQL(fileSQL = "\\Based\\" + "DataBase.db")
                result = self.CheckResult(result, labelnum = "3")
                if result != False:
                    self.AuthSucceed()
                else:
                    self.AuthFailed()
   
    def DownLoadFTP(self):
        try:
            if "temp.db" in os.listdir(self.DocDir + "\\Temp"):
                os.remove(self.DocDir + "\\Temp\\temp.db")
                ToLog("Previous temp file removed")
                
            #download from FTP to temp file
            ftp = ftplib.FTP(self.paths[0][1])
            ftp.login(self.paths[0][2], self.paths[0][3])                         
            ftp.retrbinary("RETR " + self.paths[0][4], open(self.tempName, "wb").write)
            ftp.quit()
            self.tempFile = True
        except Exception as Err:
            ToLog("Error in DownLoadFTP, Error code = " + str(Err))
            self.tempFile = False
            #raise Exception
            return False


    def CopyFromPath(self):
        try:
            if "temp.db" in os.listdir(self.DocDir + "\\Temp"):
                os.remove(self.DocDir + "\\Temp\\temp.db")
                ToLog("Previous temp file removed")

            CopyFile(self.paths[1][1], self.tempName)
            self.tempFile = True
        except Exception as Err:
            ToLog("Error in CopyFromPath, Error code = " + str(Err))
            self.tempFile = False
            #raise Exception
            return False

    def CheckSQL(self, fileSQL = "\\Temp\\" + "temp.db"):
        try:
            add = []
            All = []
            with sqlite3.connect(self.DocDir + fileSQL) as sql:
                cursor = sql.execute("SELECT RCS from SpiskiSS_Users WHERE HostName = " + "'" + self.MyName + "'")
                for row in cursor:
                    add.append(row[0])
                cursor = sql.execute("SELECT HostName, RCS from SpiskiSS_Users")
                for row in cursor:
                    All.append(row[0])
                #print("return " + str(add) + " | " + str(All))
                self.Founded = [add, All]
        except Exception as Err:
            ToLog("Error in CheckSQL with file = " + fileSQL + ", Error code = " + str(Err))

    def CheckResult(self, rcslist, labelnum):
        try:
            global Region
            rcslist = self.Founded
            #print("MyRCSLIST = " + str(rcslist))
            if not isinstance (rcslist, list):
                return False
            if len(rcslist[0]) == 0:
                ToLog("User not founded in SQL Base")
                wx.MessageBox("User " + self.MyName + " was not found in SQLBase (attemp#" + labelnum + ")", " ", wx.OK)
                return False
            elif len(rcslist[0]) > 1:
                ToLog("Multiple Users founded in SQL Base: " + str(rcslist))
                wx.MessageBox("User " + self.MyName + " was found more than once in SQLBase (attemp#" + labelnum + ")", " ", wx.OK)
                return False
            else:
                Region = rcslist[0][0]
                ToLog("Successfull authorization in attemp# " + labelnum + " with name = " + self.MyName)
        except Exception as Err:
            ToLog("Error in CheckResult, Error code = " + str(Err))
            #raise Exception
            return False

    @ExceptDecorator
    def AuthSucceed(self):
        #print("-"*40)
        #print(str(os.listdir(self.DocDir + "\\Temp")))
        if "temp.db" in os.listdir(self.DocDir + "\\Temp"):
            os.remove(self.DocDir + "\\Temp\\temp.db")
            ToLog("Previous temp file removed")
        global Region
        self.dlg.Update(4, "Успешное подключение к АС Селекторных")
        time.sleep(1)
        
    @ExceptDecorator
    def AuthFailed(self):
        if "temp.db" in os.listdir(self.DocDir + "\\Temp"):
            os.remove(self.DocDir + "\\Temp\\temp.db")
            ToLog("Previous temp file removed")
        self.dlg.Destroy()
        wx.MessageBox("Authorization failed")
        wx.Exit()
        sys.exit()
        
#=============================================
#=============================================
#=============================================
#=============================================
# stopping thread
def StopThread():
    global Thread
    try:
        if Thread == None:
            return
        elif Thread.is_alive() == False:
            #print("Thread is dead")
            return
        else:
            Thread.Stop = True
            #print("stopping thread")
            Thread.join()
            #print("thread stopped")
    except Exception as Err:
        ToLog("Ошибка при попытке остановить поток мониторинга, код ошибки = " + str(Err))
        
#=============================================
#=============================================
#=============================================
#=============================================
#SoundThread
class SoundThread(threading.Thread):
    def __init__(self, sound):
        super().__init__()
        self.sound = sound
        #print("sound = " + self.sound)

    def run(self):
        ToLog("playing melody " + self.sound)
        try:
            #playsound(self.sound)
            pygame.mixer.music.load(self.sound)
            pygame.mixer.music.play()
        except Exception as Err:
            ToLog("Failed to play melody, Error code = " + str(Err))
            #raise Exception
        
#=============================================
#=============================================
#=============================================
#=============================================
# scaling bitmaps
def ScaleBitmap(bitmap, size):
    image = bitmap.ConvertToImage()
    image = image.Scale(size[0], size[1], wx.IMAGE_QUALITY_HIGH)
    return wx.Image(image).ConvertToBitmap()

#=============================================
#=============================================
#=============================================
#=============================================
# Tolog - renew log
def ToLog(message, startThread = False):
    try:
        global LogQueue
        LogQueue.put(str(datetime.today())[10:19] + "  " + str(message) + "\n")
    except Exception as Err:
        print("Error in ToLog function, Error code = " + str(Err))
        
#=============================================
#=============================================
#=============================================
#=============================================
# Thread for saving logs
class LogThread(threading.Thread):
    def __init__(self):
        super().__init__()
        self.stop = False

    def run(self):
        global LogQueue
        ToLog("LogThread started!!!")
        self.writingQueue()
        ToLog("LogThread finished!!!")

    def writingQueue(self):
        global LogQueue, LogDir
        while True:
            try:
                if LogQueue.empty():
                    if self.stop == True:
                        print("LogThreadStopped")
                        break
                    time.sleep(1)
                    continue
                else:
                    with open(LogDir + "\\" + str(datetime.today())[0:10] + ".cfg", "a") as file:
                        while not LogQueue.empty():
                            mess = LogQueue.get_nowait()
                            file.write(mess)
                            print("Wrote to Log:\t" + mess)
                        file.close()
            except Exception as Err:
                print("Error writing to Logfile, Error code = " + str(Err))
                #raise Exception

#=============================================
#=============================================
#=============================================
#=============================================
def CopyFile(source, dist, buffer = 1024*1024):
    ToLog("CopyFile " + source + " to " + dist + "function started")
    with open(source, "rb") as SrcFile, open(dist, "wb") as DestFile:
        while True:
            copy_buffer = SrcFile.read(buffer)
            if not copy_buffer:
                break
            DestFile.write(copy_buffer)
    ToLog("CopyFile " + source + " to " + dist + "function finished")

#=============================================
#=============================================
#=============================================
#=============================================
# CreateFolders
def CreateFolders():
    global MyDir, MonitOpen
    MyDir = os.getcwd()
    MonitOpen = False
    
    try:
        os.mkdir(MyDir + "\\Conferense_Lists")
    except Exception as Err:
        print("Error creating folder of Lists, Error code = " + str(Err))

    try:
        os.mkdir(MyDir + "\\Logs")
    except Exception as Err:
        print("Error creating folder of Logs, Error code = " + str(Err))

    try:
        os.mkdir(MyDir + "\\Based")
    except Exception as Err:
        print("Error creating folder of Bases, Error code = " + str(Err))

    #removing tempfile ftp
    try:
        os.remove(MyDir + "\\temp.bd")
    except Exception as Err:
        print("Error deleting temp file, Error code = " + str(Err))
#=============================================
#=============================================
#=============================================
#=============================================
# OpenNon-LocalBase
def OpenPaths(DocDir):
    try:
        Paths = []
        with sqlite3.connect(DocDir + "\\Based\\" + "DataBase.db") as conn:
            cursor = conn.execute("SELECT * from Paths")
            for row in cursor:
                Paths.append(row)
            #conn.close()
        #return Paths
    except Exception as Err:
        ToLog("Error in OpenPaths, Error code = " + str(Err))
        #raise Exception
        sys.exit()
    else:
        ToLog("OpenPaths successed, paths are = " + str(Paths))
        return Paths

#=============================================
#=============================================
#=============================================
# OpenNon-LocalBase
def LoadColours(DocDir):
    global TimeRefresh
    try:
        with open(DocDir + "\\Based\\" + "colours.txt", "r") as file:
            info = file.read().splitlines()
            file.close()
        #print("colours = " + str(info))
        colours = ["#bfdee8", "#EBF7FF", "#EBF777"]
        TimeRefresh = 30
        for word in info:
            if len(word.strip()) > 0:
                #word = word.split()[:]
                #print("word = " + str(word))
                if word.find("colourback=") != -1:
                    #print("find colourback=")
                    colours[0] = word.split()[1][:]
                elif word.find("colourbtn=") != -1:
                    #print("find colourbtn=")
                    colours[1] = word.split()[1][:]
                elif word.find("#colourbtnmonit=") != -1:
                    #print("find colourbtnmonit=")
                    colours[2] = word.split()[1][:]
                elif word.find("#timerefresh=") != -1:
                    TimeRefresh = int(word.split()[1][:])
                    if TimeRefresh < 5:
                        TimeRefresh = 5
        #print("coloursafter func = " + str(colours) + ", time Refresh = " + str(TimeRefresh))
        ToLog("Загружены цвета для приложения = " + str(colours))
        ToLog("Загружено время обновления мониторинга = " + str(TimeRefresh))
                                       
    except Exception as Err:
        ToLog("Error reading ColoursandTime, Error code = " + str(Err))
        TimeRefresh = 30
        colours = ["#bfdee8", "#EBF7FF", "#EBF777"]
        #raise Exception

    return colours
        
#=============================================
#=============================================
#=============================================
#=============================================
# ClearOldLogs
def ClearLogs():
    global LogDir
    try:
        while len(os.listdir(LogDir)) >= 10:
            if len(os.listdir(LogDir)) < 10:
                    break
            try:
                os.remove(os.path.abspath(FindOldest(LogDir)))
                ToLog("DELETING FILE " + str(FindOldest(LogDir)))
            except Exception as Err:
                ToLog("Old file with logs wasn't deleted, Error code = " + str(Err))
                #raise Exception
                break
    except Exception as Err:
        ToLog("Error of clearing dir with logs, Error code = " + str(Err))
        #raise Exception
#=============================================
#=============================================
#=============================================
#=============================================   
# DeleteOldest
def FindOldest(Dir):
    try:
        List = os.listdir(Dir)
        fullPath = [Dir + "/{0}".format(x) for x in List]
        oldestFile = min(fullPath, key = os.path.getctime)
        return oldestFile
    except Exception as Err:
        ToLog("Error of finding oldest file in dir, Error code = " + str(Err))
        #raise Exception
        return False
#=============================================
#=============================================
#=============================================
#=============================================
def FindMyDir(nameDir, subDirs = None):
    try:
        if "Documents" in os.listdir(os.path.expanduser("~")):
            DocDir = os.path.expanduser("~") + "\\Documents"
        else:
            os.mkdir(os.path.expanduser("~") + "\\Documents")
            DocDir = os.path.expanduser("~") + "\\Documents"
        if nameDir not in os.listdir(DocDir):
            os.mkdir(DocDir + "\\" + nameDir)
            ToLog(nameDir + "folder was Created")

        DocDir = DocDir + "\\" + nameDir
        if isinstance (subDirs, list):
            for word in subDirs:
                if word not in os.listdir(DocDir):
                    os.mkdir(DocDir + "\\" + word)
                    ToLog(word + " folder was Created")
        if "sound.wav" not in os.listdir(DocDir + "\\Sounds"):
            CopyFile(os.getcwd() + "\\Sounds\\sound.wav", DocDir + "\\Sounds\\sound.wav")
        if "colours.txt" not in os.listdir(DocDir + "\\Based"):
            CopyFile(os.getcwd() + "\\Based\\colours.txt", DocDir + "\\Based\\colours.txt")
        if "DataBase.db" not in os.listdir(DocDir + "\\Based"):
            CopyFile(os.getcwd() + "\\Based\\DataBase.db", DocDir + "\\Based\\DataBase.db")
    except Exception as Err:
        ToLog("Error in FindMyDir, Error code = " + str(Err))  
        #raise Exception
        return os.path.expanduser("~") + "\\" + nameDir
    else:
        ToLog("FindMyDir finished succesfully")
        return DocDir

#=============================================
#=============================================
#=============================================
#=============================================
# Определение локали!
locale.setlocale(locale.LC_ALL, "")

global LogDir, LogQueue, MyDate, threadLog, MonitOpen
LogQueue = queue.Queue()
MyDate = "18.04.2024"
MonitOpen = False

ToLog("\n\n" + "!" * 40)
ToLog("Application started")

DocDir = FindMyDir(nameDir = "SpiskiSSFiles", subDirs = ["Logs", "Conference_Lists", "Based", "Sounds", "Temp"])
LogDir = DocDir + "\\Logs"
ClearLogs()

threadLog = LogThread()
threadLog.setDaemon(True)
threadLog.start()

Paths = OpenPaths(DocDir = DocDir)

ex = wx.App()

AuthDlg(DocDir, Paths)

ex.MainLoop()





